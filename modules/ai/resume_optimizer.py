import os
import json
import re
from typing import List, Tuple
from openai import OpenAI
from pathlib import Path
import pdfplumber
from PyPDF2 import PdfReader, PdfWriter
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from io import BytesIO
import textwrap
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.platypus import Paragraph, SimpleDocTemplate
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
from docx.shared import Pt
import docx2pdf
import pypandoc
from config.secrets import llm_api_key  # Import the API key from secrets.py

class ResumeOptimizer:
    def __init__(self, api_key: str):
        self.client = OpenAI(api_key=api_key)

    def extract_keywords(self, job_description: str) -> List[str]:
        """Extract top 15 keywords from job description using OpenAI API."""
        try:
            response = self.client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {
                        "role": "system",
                        "content": (
                            "You are an ATS expert that analyzes job descriptions and "
                            "identifies the most important keywords for ATS systems like Greenhouse and Workday. "
                            "Return exactly 15 most important keywords as a JSON object in this format:\n"
                            '{"keywords": ["keyword1", "keyword2", ..., "keyword15"]}'
                        )
                    },
                    {
                        "role": "user",
                        "content": (
                            "Analyze this job description and identify the top 15 most important keywords "
                            "that should be included in a resume to pass ATS systems:\n\n" + job_description
                        )
                    }
                ]
            )

            content = response.choices[0].message.content.strip()
            # print(f"\n🔍 Raw GPT response:\n{content}\n")

            # Try parsing directly
            try:
                result = json.loads(content)
            except json.JSONDecodeError:
                # Try extracting JSON substring
                match = re.search(r'\{.*\}', content, re.DOTALL)
                if match:
                    result = json.loads(match.group())
                else:
                    raise ValueError("No valid JSON found in response.")

            # print(f"✅ Extracted keywords: {result}")
            return result.get('keywords', [])

        except Exception as e:
            print(f"❌ Error extracting keywords: {e}")
            return []
        
    def read_pdf_resume(self, resume_path: str) -> Tuple[str, List[str], dict]:
        """Read PDF resume and extract the entire skills section with formatting information."""
        try:
            with pdfplumber.open(resume_path) as pdf:
                content = ""
                skills_page = 0
                skills_bbox = None
                current_skills = []

                for page_num, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    content += page_text + "\n"

                    lines = page_text.splitlines()
                    skills_started = False
                    skills_lines = []

                    for i, line in enumerate(lines):
                        if not skills_started and re.match(r"(?i)^skills\s*:", line.strip()):
                            skills_started = True
                            skills_lines.append(line)
                            continue
                        if skills_started:
                            # Stop if we hit the next section (all caps heading or known section)
                            if re.match(r"^[A-Z\s]{5,}$", line.strip()):
                                break
                            skills_lines.append(line)

                    if skills_lines:
                        skills_page = page_num
                        skills_text = " ".join(skills_lines)
                        # Extract the part after "Skills:"
                        skills_raw = re.sub(r"(?i)^skills\s*:\s*", "", skills_text)
                        current_skills = [skill.strip() for skill in skills_raw.split(',') if skill.strip()]

                        # Extract skills bbox
                        words = page.extract_words()
                        for word in words:
                            if "skills" in word['text'].lower():
                                skills_bbox = {
                                    'x0': word['x0'],
                                    'x1': word['x1'],
                                    'top': word['top'],
                                    'bottom': word['bottom']
                                }
                                break

                        break  # We found the section, no need to check other pages

                if not current_skills:
                    raise ValueError("Skills section not found in resume")

                format_info = {
                    'skills_page': skills_page,
                    'skills_bbox': skills_bbox
                }

                # print(f"Skills section found on page {skills_page + 1} with bbox: {skills_bbox}")
                # print(f"Extracted skills: {current_skills}")
                return content, current_skills, format_info

        except Exception as e:
            print(f"Error reading PDF resume: {e}")
            return "", [], {}
        
    def optimize_skills(self, current_skills: List[str], new_keywords: List[str], job_description: str) -> List[str]:
        """Optimize skills list by incorporating new keywords while maintaining the same length."""
        current_set = set(current_skills)
        new_set = set(new_keywords)
        # print(f"Current skills: {current_set}")
        # print(f"New keywords: {new_set}")
        
        # Keywords to add (that aren't already present)
        to_add = new_set - current_set
        # print(f"New keywords to add: {to_add}")
        
        if not to_add:
            return current_skills
            
        # Calculate how many skills we need to remove
        num_to_remove = len(to_add)
        
        # Remove least relevant skills (those not in new_keywords)
        skills_to_keep = list(current_set - to_add)[:len(current_skills) - num_to_remove]
        
        # Combine kept skills with new keywords
        optimized_skills = skills_to_keep + list(to_add)

        # Order the optimized skills with the best order for ATS systems optimized for the job description using OpenAI API
        # try:
        #     response = self.client.chat.completions.create(
        #     model="gpt-4o",
        #     messages=[
        #         {
        #         "role": "system",
        #         "content": (
        #             "You are an ATS optimization expert. Given a list of skills and a job description, "
        #             "reorder the skills to maximize relevance and impact for ATS systems. Make sure to always include all the skills from the original list. "
        #             "Return the reordered skills as a JSON object in this format without the json code block:\n"
        #             '{"ordered_skills": ["skill1", "skill2", ..., "skillN"]}'
        #         )
        #         },
        #         {
        #         "role": "user",
        #         "content": (
        #             f"Here is the job description:\n{job_description}\n\n"
        #             f"Here is the list of skills to reorder:\n{optimized_skills}"
        #         )
        #         }
        #     ]
        #     )

        #     content = response.choices[0].message.content.strip()
        #     print(f"\n🔍 Raw GPT response:\n{content}\n")
        #     # Try parsing the response
        #     try:
        #         result = json.loads(content)
        #         ordered_skills = result.get('ordered_skills', optimized_skills)
        #         if ordered_skills and isinstance(ordered_skills, list):
        #             optimized_skills = ordered_skills
        #     except json.JSONDecodeError:
        #         print("❌ Error parsing OpenAI response for skill ordering. Using original order.")

        # except Exception as e:
        #     print(f"❌ Error reordering skills using OpenAI API: {e}")
        
        return optimized_skills[:len(current_skills)]
    
    def save_optimized_pdf(self, resume_path: str, optimized_skills: List[str], format_info: dict, output_path: str):
        try:
            # Step 1: Use pdfplumber to extract structured text
            with pdfplumber.open(resume_path) as pdf:
                full_text_lines = []
                skills_section_found = False
                skills_started = False

                for page in pdf.pages:
                    page_lines = page.extract_text().splitlines()
                    i = 0
                    while i < len(page_lines):
                        line = page_lines[i]
                        if not skills_started and re.match(r"(?i)^skills\s*:", line.strip()):
                            skills_started = True
                            skills_section_found = True

                            # Format the new skills section
                            wrapped_skills = textwrap.wrap(", ".join(optimized_skills), width=110)
                            wrapped_skills[0] = "Skills: " + wrapped_skills[0]  # Add label to first line
                            full_text_lines.extend(wrapped_skills)

                            # Skip original skills block
                            i += 1
                            while i < len(page_lines):
                                next_line = page_lines[i]
                                if re.match(r"^[A-Z\s]{5,}$", next_line.strip()):
                                    break
                                i += 1
                            continue
                        if skills_started and re.match(r"^[A-Z\s]{5,}$", line.strip()):
                            skills_started = False
                        full_text_lines.append(line)
                        i += 1

                if not skills_section_found:
                    raise ValueError("Skills section not found in the resume.")

            # Step 2: Rebuild PDF using ReportLab
            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter,
                                    rightMargin=40, leftMargin=40,
                                    topMargin=60, bottomMargin=40)
            styles = getSampleStyleSheet()
            story = []

            for para in "\n".join(full_text_lines).split('\n\n'):
                for line in para.strip().splitlines():
                    story.append(Paragraph(line.strip(), styles["Normal"]))

            doc.build(story)

            # Step 3: Save to disk
            with open(output_path, 'wb') as f:
                f.write(buffer.getvalue())

            print(f"✅ Successfully saved updated PDF to: {output_path}")

        except Exception as e:
            print(f"❌ Error saving optimized PDF resume: {e}")

    def read_docx_resume(self, resume_path: str) -> Tuple[Document, List[str]]:
        """Read a DOCX resume and extract the skills section."""
        try:
            doc = Document(resume_path)
            skills_text = ""
            skills_started = False
            skills_lines = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if not skills_started and re.match(r"(?i)^skills\s*:", text):
                    skills_started = True
                    skills_lines.append(text)
                    continue
                if skills_started:
                    if re.match(r"^[A-Z\s]{5,}$", text):
                        break
                    skills_lines.append(text)

            if not skills_lines:
                raise ValueError("Skills section not found in DOCX resume")

            skills_text = " ".join(skills_lines)
            skills_raw = re.sub(r"(?i)^skills\s*:\s*", "", skills_text)
            current_skills = [s.strip() for s in skills_raw.split(',') if s.strip()]

            return doc, current_skills

        except Exception as e:
            print(f"Error reading DOCX resume: {e}")
            return None, []

    def save_optimized_docx(self, doc: Document, optimized_skills: List[str], output_path: str):
        """Update the skills section in the DOCX and save it."""
        try:
            skills_started = False
            skills_index = None
            end_index = None

            # Identify where the Skills section starts and ends
            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if not skills_started and re.match(r"(?i)^skills\s*:", text):
                    skills_started = True
                    skills_index = i
                    continue
                if skills_started and re.match(r"^[A-Z\s]{5,}$", text):
                    end_index = i
                    break

            if skills_index is None:
                raise ValueError("Skills section not found in DOCX resume")

            if end_index is None:
                end_index = len(doc.paragraphs)

            # Delete only paragraphs after the SKILLS header and before next section
            for i in range(end_index - 1, skills_index, -1):
                p = doc.paragraphs[i]
                p._element.getparent().remove(p._element)

            # Replace the content of the skills paragraph with the first line
            skills_text_wrapped = textwrap.wrap(", ".join(optimized_skills), width=120)
            doc.paragraphs[skills_index].text = "Skills: " + skills_text_wrapped[0]

            # Add the rest of the lines after the current paragraph
            insert_after = doc.paragraphs[skills_index]._element
            for line in skills_text_wrapped[1:]:
                new_para = doc.add_paragraph(line)
                new_para.style = doc.paragraphs[skills_index].style
                insert_after.addnext(new_para._element)
                insert_after = new_para._element

            doc.save(output_path)
            print(f"✅ Successfully saved updated DOCX resume to: {output_path}")

        except Exception as e:
            print(f"❌ Error saving optimized DOCX resume: {e}")

    # Convert docx to pdf
    def convert_docx_to_pdf(self, docx_path: str, pdf_path: str):
        try:
            # pypandoc.convert_file(docx_path, 'pdf', outputfile=pdf_path)
            # docx2pdf.convert(docx_path, pdf_path)
            # Use LibreOffice to convert DOCX to PDF
            import subprocess
            libreoffice_path = "/Applications/LibreOffice.app/Contents/MacOS/soffice"
            subprocess.run([libreoffice_path, '--headless', '--convert-to', 'pdf', docx_path, '--outdir', str(Path(pdf_path).parent)], check=True)
            print(f"✅ Successfully converted DOCX to PDF: {pdf_path}")
        except ImportError:
            print("❌ pypandoc module not found. Please install it using 'pip install pypandoc'.")
        except Exception as e:
            print(f"❌ Error converting DOCX to PDF: {e}")

# Function to run the resume optimization process given job description and resume path similar to main()
def run_resume_optimization(job_description: str, resume_path: str, job_id: str):
    # Get OpenAI API key
    api_key = llm_api_key
    if not api_key:
        print("Please set the OPENAI_API_KEY environment variable")
        return

    optimizer = ResumeOptimizer(api_key)

    # Extract keywords from job description
    print("\nExtracting keywords from job description...")
    keywords = optimizer.extract_keywords(job_description)
    print(f"Top keywords identified: {', '.join(keywords)}")

    # Check if resume file exists
    if not Path(resume_path).exists():
        print("Resume file not found!")
        return

    ext = Path(resume_path).suffix.lower()
    output_path = None

    # Process resume based on file type
    print("\nProcessing resume...")
    if ext == ".pdf":
        content, current_skills, format_info = optimizer.read_pdf_resume(resume_path)
        if not content or not current_skills:
            return
        optimized_skills = optimizer.optimize_skills(current_skills, keywords, job_description)
        output_path = Path(resume_path).parent.parent / "temp" / f"danny_basavaraju_{job_id}.pdf"
        optimizer.save_optimized_pdf(resume_path, optimized_skills, format_info, str(output_path))

    elif ext == ".docx":
        doc, current_skills = optimizer.read_docx_resume(resume_path)
        if not doc or not current_skills:
            return
        optimized_skills = optimizer.optimize_skills(current_skills, keywords, job_description)
        output_path_docx = Path(resume_path).parent.parent / "temp" / f"danny_basavaraju_{job_id}.docx"
        output_path_pdf = Path(resume_path).parent.parent / "temp" / f"danny_basavaraju_{job_id}.pdf"
        optimizer.save_optimized_docx(doc, optimized_skills, str(output_path_docx))
        optimizer.convert_docx_to_pdf(str(output_path_docx), str(output_path_pdf))
        output_path = output_path_pdf

    else:
        print("❌ Unsupported file format. Please use PDF or DOCX.")
        return

    print(f"\n🎯 Optimized resume saved to: {output_path}")
    return output_path

def main():
    # Get OpenAI API key from environment variable
    api_key = llm_api_key
    if not api_key:
        print("Please set the OPENAI_API_KEY environment variable")
        return

    optimizer = ResumeOptimizer(api_key)

    # Get job description
    print("Please paste the job description (press Ctrl+D when finished):")
    job_description = ""
    try:
        while True:
            line = input()
            job_description += line + "\n"
    except EOFError:
        pass

    # Get resume path
    resume_path = input("\nEnter the path to your resume (PDF or DOCX): ").strip()
    if not Path(resume_path).exists():
        print("Resume file not found!")
        return
    
    ext = Path(resume_path).suffix.lower()


    # Extract keywords from job description
    print("\nExtracting keywords from job description...")
    keywords = optimizer.extract_keywords(job_description)
    print(f"Top keywords identified: {', '.join(keywords)}")

    # Read and process resume
    print("\nProcessing resume...")
    output_path = None

    if ext == ".pdf":
        content, current_skills, format_info = optimizer.read_pdf_resume(resume_path)
        if not content or not current_skills:
            return
        optimized_skills = optimizer.optimize_skills(current_skills, keywords, job_description)
        output_path = Path(resume_path).with_name(f"optimized_{Path(resume_path).stem}.pdf")
        optimizer.save_optimized_pdf(resume_path, optimized_skills, format_info, str(output_path))

    elif ext == ".docx":
        doc, current_skills = optimizer.read_docx_resume(resume_path)
        if not doc or not current_skills:
            return
        optimized_skills = optimizer.optimize_skills(current_skills, keywords, job_description)
        output_path_docx = Path(resume_path).with_name(f"optimized_{Path(resume_path).stem}.docx")
        output_path_pdf = Path(resume_path).with_name(f"optimized_{Path(resume_path).stem}.pdf")
        optimizer.save_optimized_docx(doc, optimized_skills, str(output_path_docx))
        optimizer.convert_docx_to_pdf(str(output_path_docx), str(output_path_pdf))
        output_path = output_path_pdf

    else:
        print("❌ Unsupported file format. Please use PDF or DOCX.")
        return

    print(f"\n🎯 Optimized resume saved to: {output_path}")

# if __name__ == "__main__":
#     main()
